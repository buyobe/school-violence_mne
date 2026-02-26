import io, base64, json
from django.shortcuts import render
from django.http import HttpResponse
from django.db.models import Count
from data_collection.models import Student, Teacher, Parent
from django.conf import settings
import os, datetime
from django.contrib.auth.decorators import login_required


# ReportLab for PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

# Word export
from docx import Document
from docx.shared import Inches

# Excel export
import pandas as pd




@login_required
def dashboard(request):
    total_students = Student.objects.count()
    total_teachers = Teacher.objects.count()
    total_parents = Parent.objects.count()

    # Example: reporting rate
    reporting_students = Student.objects.filter(reporting_violence=True).count()
    student_reporting_rate = round((reporting_students / total_students * 100), 1) if total_students else 0

    combined_chart = {
        "labels": json.dumps(["Students", "Teachers", "Parents"]),
        "values": json.dumps([
            reporting_students,
            Teacher.objects.filter(reporting_violence=True).count(),
            Parent.objects.filter(reporting_violence=True).count()
        ])
    }

    context = {
        "total_students": total_students,
        "total_teachers": total_teachers,
        "total_parents": total_parents,
        "student_reporting_rate": student_reporting_rate,
        "combined_chart": combined_chart,
    }
    return render(request, "visualization/dashboard.html", context)

@login_required
def demographics_view(request):
    # Gender distribution (students only for now)
    gender_distribution = Student.objects.values("gender").annotate(count=Count("id"))

    # Age group distribution
    age_group_distribution = Student.objects.values("age_group").annotate(count=Count("id"))

    # Disability status
    disability_distribution = Student.objects.values("disability_status").annotate(count=Count("id"))

    # Role-based cases (students, teachers, parents)
    role_distribution = [
        {"role": "Students", "count": Student.objects.count()},
        {"role": "Teachers", "count": Teacher.objects.count()},
        {"role": "Parents", "count": Parent.objects.count()},
    ]

    return render(request, "visualization/demographics.html", {
        "gender_distribution": gender_distribution,
        "age_group_distribution": age_group_distribution,
        "disability_distribution": disability_distribution,
        "role_distribution": role_distribution,
    })

@login_required
def trends_view(request):
    category = request.GET.get("category", "region")

    # Get field mapping for each role
    field_map = {
        "region": {"student": "region", "teacher": "region", "parent": "region"},
        "district": {"student": "district", "teacher": "district", "parent": "district"},
        "school": {"student": "school", "teacher": "school", "parent": "school"},
        "gender": {"student": "gender", "teacher": "gender", "parent": "gender"},
        "age_group": {"student": "age_group", "teacher": "age_group", "parent": "age_group"},
        "disability": {"student": "disability_status", "teacher": None, "parent": None},
        "violence_type": {"student": "forms_of_violence", "teacher": "forms_of_violence", "parent": "forms_of_violence"},
        "perpetrator": {"student": "perpetrators", "teacher": None, "parent": None},
        "reporting": {"student": "reporting_violence", "teacher": "reporting_violence", "parent": "reporting_violence"},
        "system_effectiveness": {
            "student": "effectiveness_reporting_system",
            "teacher": "effective_handling_vac",
            "parent": "effectiveness_positive_punishment",
        },
    }

    fields = field_map.get(category)

    # Query each role only if the field exists
    student_data = Student.objects.values(fields["student"]).annotate(count=Count("id")).order_by(fields["student"]) if fields["student"] else []
    teacher_data = Teacher.objects.values(fields["teacher"]).annotate(count=Count("id")).order_by(fields["teacher"]) if fields["teacher"] else []
    parent_data = Parent.objects.values(fields["parent"]).annotate(count=Count("id")).order_by(fields["parent"]) if fields["parent"] else []

    # Labels from student dataset (fallback to teacher/parent if empty)
    labels_source = student_data or teacher_data or parent_data
    labels = [item[fields["student"] or fields["teacher"] or fields["parent"]] or "Unknown" for item in labels_source]

    # Counts aligned to labels
    def align_counts(data, field):
        return [next((d["count"] for d in data if d[field] == label), 0) for label in labels] if field else [0]*len(labels)

    student_counts = align_counts(student_data, fields["student"])
    teacher_counts = align_counts(teacher_data, fields["teacher"])
    parent_counts = align_counts(parent_data, fields["parent"])

    return render(request, "visualization/trends.html", {
        "selected_category": category,
        "labels": labels,
        "student_counts": student_counts,
        "teacher_counts": teacher_counts,
        "parent_counts": parent_counts,
    })




    # Helper to map counts
    @login_required
    def get_counts(qs):
        counts = {label: 0 for label in labels}
        for entry in qs:
            if field == "disability_status":
                label = "With Disability" if entry[field] else "Without Disability"
            else:
                label = entry[field] if entry[field] else "Unknown"
            counts[label] = entry["count"]
        return [counts[label] for label in labels]

    student_counts = get_counts(student_qs)
    teacher_counts = get_counts(teacher_qs)
    parent_counts = get_counts(parent_qs)

    context = {
        "selected_category": category,
        "labels": json.dumps(labels),
        "student_counts": json.dumps(student_counts),
        "teacher_counts": json.dumps(teacher_counts),
        "parent_counts": json.dumps(parent_counts),
    }
    return render(request, "visualization/trends.html", context)



@login_required
def reports_view(request):
    # Aggregations
    top_region = Student.objects.values("region").annotate(count=Count("id")).order_by("-count").first()
    top_school = Student.objects.values("school").annotate(count=Count("id")).order_by("-count").first()
    top_violence = Student.objects.values("forms_of_violence").annotate(count=Count("id")).order_by("-count").first()
    reporting_effectiveness = Student.objects.values("effectiveness_reporting_system").annotate(count=Count("id"))

    # Executive summary text
    executive_summary = "This report highlights key findings from the school violence monitoring system. "
    if top_region:
        executive_summary += f"The region most affected is {top_region['region']} with {top_region['count']} reported cases. "
    if top_school:
        executive_summary += f"The school with the highest cases is {top_school['school']} ({top_school['count']} cases). "
    if top_violence:
        executive_summary += f"The most common form of violence is {top_violence['forms_of_violence']} ({top_violence['count']} cases). "
    if reporting_effectiveness:
        effective = sum([r["count"] for r in reporting_effectiveness if r["effectiveness_reporting_system"] == "Effective"])
        ineffective = sum([r["count"] for r in reporting_effectiveness if r["effectiveness_reporting_system"] == "Ineffective"])
        executive_summary += f"Reporting systems were rated effective in {effective} cases and ineffective in {ineffective} cases. "
    else:
        effective, ineffective = 0, 0

    # Recommendations
    recommendations = []
    if ineffective > effective:
        recommendations.append("Strengthen confidential reporting channels and child protection committees.")
    if top_violence and top_violence["forms_of_violence"].lower() == "corporal punishment":
        recommendations.append("Expand teacher training on positive discipline methods.")
    if top_region and top_region["count"] > 100:
        recommendations.append(f"Prioritize interventions in {top_region['region']}.")


    # Violence type distribution (students only for now)
    violence_qs = Student.objects.values("forms_of_violence").annotate(count=Count("id"))
    violence_labels = [v["forms_of_violence"] or "Unknown" for v in violence_qs]
    violence_counts = [v["count"] for v in violence_qs]

    # Perpetrator distribution
    perpetrator_qs = Student.objects.values("perpetrators").annotate(count=Count("id"))
    perpetrator_labels = [p["perpetrators"] or "Unknown" for p in perpetrator_qs]
    perpetrator_counts = [p["count"] for p in perpetrator_qs]

    # Reporting effectiveness
    reporting_qs = Student.objects.values("effectiveness_reporting_system").annotate(count=Count("id"))
    reporting_labels = [r["effectiveness_reporting_system"] or "Unknown" for r in reporting_qs]
    reporting_counts = [r["count"] for r in reporting_qs]

    context = {
        "executive_summary": executive_summary,
        "recommendations": recommendations,
        "violence_labels": json.dumps(violence_labels),
        "violence_counts": json.dumps(violence_counts),
        "perpetrator_labels": json.dumps(perpetrator_labels),
        "perpetrator_counts": json.dumps(perpetrator_counts),
        "reporting_labels": json.dumps(reporting_labels),
        "reporting_counts": json.dumps(reporting_counts),
    }
    return render(request, "visualization/reports.html", context)


@login_required
def export_pdf(request):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    # --- Cover Page ---
    logo_path = os.path.join(settings.BASE_DIR, "static", "images", "fawe_logo.png")
    try:
        story.append(Image(logo_path, width=120, height=120))
    except Exception:
        story.append(Paragraph("NGO Logo Missing", styles['Normal']))
    story.append(Spacer(1, 24))
    story.append(Paragraph("School Violence Monitoring & Evaluation Report", styles['Title']))
    story.append(Spacer(1, 36))
    story.append(Paragraph("Prepared by: FAWE TANZANIA", styles['Normal']))
    story.append(Paragraph("Date: " + str(datetime.date.today()), styles['Normal']))
    story.append(Spacer(1, 48))

    # --- Executive Summary (on same page as cover info) ---
    story.append(Paragraph("Executive Summary", styles['Heading1']))
    story.append(Paragraph(request.POST.get("executive_summary", "No summary available"), styles['Normal']))
    story.append(Spacer(1, 24))

    # ✅ Now insert a page break before recommendations/charts
    story.append(PageBreak())

    # --- Recommendations ---
    recs = request.POST.getlist("recommendations")
    if recs:
        story.append(Paragraph("Recommendations", styles['Heading1']))
        for rec in recs:
            story.append(Paragraph(f"- {rec}", styles['Normal']))
        story.append(Spacer(1, 24))

    # --- Charts ---
    for chart_key in ["violence_chart", "perpetrator_chart", "reporting_chart"]:
        chart_data = request.POST.get(chart_key)
        if chart_data:
            chart_bytes = base64.b64decode(chart_data.split(",")[1])
            chart_buffer = io.BytesIO(chart_bytes)
            story.append(Image(chart_buffer, width=400, height=250))
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="ngo_report.pdf"'
    return response


@login_required
def export_word(request):
    doc = Document()

    # --- Cover Page ---
    doc.add_picture("static/images/fawe_logo.png", width=Inches(2))  # adjust path
    doc.add_paragraph("School Violence Monitoring & Evaluation Report", style="Title")
    doc.add_paragraph("Prepared by: FAWE TANZANIA")
    doc.add_paragraph("Date: " + str(datetime.date.today()))
    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(request.POST.get("executive_summary", "No summary available"))

    # --- Recommendations ---
    recs = request.POST.getlist("recommendations")
    if recs:
        doc.add_heading("Recommendations", level=1)
        for rec in recs:
            doc.add_paragraph(rec, style="ListBullet")

    # --- Charts ---
    for chart_key in ["violence_chart", "perpetrator_chart", "reporting_chart"]:
        chart_data = request.POST.get(chart_key)
        if chart_data:
            chart_bytes = base64.b64decode(chart_data.split(",")[1])
            chart_buffer = io.BytesIO(chart_bytes)
            doc.add_picture(chart_buffer, width=Inches(5))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="ngo_report.docx"'
    return response


@login_required
def export_excel(request):
    data = {
        "Executive Summary": [request.GET.get("executive_summary", "No summary available")],
        "Recommendations": request.GET.getlist("recommendations"),
    }
    df = pd.DataFrame(data)

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Report")

    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = 'attachment; filename="ngo_report.xlsx"'
    return response



